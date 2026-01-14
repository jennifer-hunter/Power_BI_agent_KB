"""
Directory Flattener for Power BI Projects

Recursively reads all files from a directory and outputs them to a Word document
in a structured format that can be parsed back to restore the original files.
"""

import os
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK


# Load environment variables
load_dotenv()

# File delimiter markers (used for parsing back)
FILE_START_MARKER = "═══ FILE: {path} ═══"
FILE_END_MARKER = "═══ END FILE ═══"

# Binary/unreadable file extensions to skip
BINARY_EXTENSIONS = {
    '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.webp',
    '.pdf', '.zip', '.tar', '.gz', '.rar', '.7z',
    '.exe', '.dll', '.so', '.dylib',
    '.mp3', '.mp4', '.wav', '.avi', '.mov',
    '.pyc', '.pyo', '.class',
}


def is_binary_file(file_path: Path) -> bool:
    """Check if a file is likely binary based on extension or content."""
    if file_path.suffix.lower() in BINARY_EXTENSIONS:
        return True

    # Try to read a small chunk to detect binary content
    try:
        with open(file_path, 'rb') as f:
            chunk = f.read(1024)
            if b'\x00' in chunk:  # Null bytes indicate binary
                return True
    except Exception:
        return True

    return False


def read_file_content(file_path: Path) -> str | None:
    """Read file content, trying different encodings."""
    encodings = ['utf-8', 'utf-8-sig', 'utf-16', 'latin-1']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as e:
            print(f"  Error reading {file_path}: {e}")
            return None

    return None


def flatten_directory_to_word(input_dir: str, output_file: str) -> None:
    """
    Flatten a directory's contents into a Word document.

    Args:
        input_dir: Path to the directory to flatten
        output_file: Path to the output Word document
    """
    input_path = Path(input_dir)

    if not input_path.exists():
        raise FileNotFoundError(f"Input directory not found: {input_dir}")

    if not input_path.is_dir():
        raise NotADirectoryError(f"Input path is not a directory: {input_dir}")

    # Create Word document
    doc = Document()

    # Add title
    title = doc.add_heading(f"Directory Contents: {input_path.name}", level=0)

    # Add metadata paragraph
    meta = doc.add_paragraph()
    meta.add_run(f"Source: {input_path.absolute()}\n").bold = True
    meta.add_run("Format: Each file is wrapped in markers for restoration.\n")
    meta.add_run("Do not modify the FILE markers - only edit content between them.\n").italic = True

    doc.add_paragraph()  # Spacer

    # Collect all files recursively
    files_processed = 0
    files_skipped = 0

    for file_path in sorted(input_path.rglob('*')):
        if file_path.is_dir():
            continue

        # Get relative path for the marker
        relative_path = file_path.relative_to(input_path)

        # Skip binary files
        if is_binary_file(file_path):
            print(f"  Skipping binary: {relative_path}")
            files_skipped += 1
            continue

        # Read file content
        content = read_file_content(file_path)
        if content is None:
            print(f"  Could not read: {relative_path}")
            files_skipped += 1
            continue

        # Add file start marker as heading
        start_marker = FILE_START_MARKER.format(path=relative_path.as_posix())
        doc.add_heading(start_marker, level=2)

        # Add file content in a code-style paragraph
        content_para = doc.add_paragraph()
        content_run = content_para.add_run(content)
        content_run.font.name = 'Consolas'
        content_run.font.size = Pt(9)

        # Add end marker
        end_para = doc.add_paragraph()
        end_run = end_para.add_run(FILE_END_MARKER)
        end_run.bold = True

        # Add page break between files for readability
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        files_processed += 1
        print(f"  Added: {relative_path}")

    # Save document
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print(f"\n{'='*50}")
    print(f"Flattening complete!")
    print(f"  Files processed: {files_processed}")
    print(f"  Files skipped: {files_skipped}")
    print(f"  Output: {output_path.absolute()}")


def main():
    """Main entry point."""
    input_dir = os.getenv('INPUT_DIR')
    output_file = os.getenv('OUTPUT_FILE')

    if not input_dir:
        raise ValueError("INPUT_DIR not set in .env file")
    if not output_file:
        raise ValueError("OUTPUT_FILE not set in .env file")

    print(f"Directory Flattener")
    print(f"{'='*50}")
    print(f"Input:  {input_dir}")
    print(f"Output: {output_file}")
    print(f"{'='*50}\n")

    flatten_directory_to_word(input_dir, output_file)


if __name__ == '__main__':
    main()
